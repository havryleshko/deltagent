from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from agent.models import AgentRun

_SECTION_TITLES = frozenset(
    {"EXECUTIVE SUMMARY", "LINE COMMENTARY", "INSIGNIFICANT VARIANCES"}
)


def _slug_period(period: str) -> str:
    raw = (period or "").strip()
    cleaned = re.sub(r"[^\w\s-]", "", raw, flags=re.UNICODE)
    slug = re.sub(r"\s+", "_", cleaned.strip())
    return slug or "unknown_period"


def build_export_basename(period: str, extension: str) -> str:
    ext = extension.lstrip(".").lower()
    return f"variance_commentary_{_slug_period(period)}.{ext}"


def render_run_markdown(agent_run: AgentRun) -> str:
    lines: list[str] = ["EXECUTIVE SUMMARY", agent_run.executive_summary.strip(), "", "LINE COMMENTARY", ""]
    for item in agent_run.line_items:
        if item.review_status not in {"accepted", "edited"}:
            continue
        lines.append(item.header)
        lines.append(item.final_commentary.strip())
        if item.sources:
            lines.append("")
            lines.append("Sources")
            for source in item.sources:
                lines.append(
                    f"- {source.source_type} - {source.timestamp} - {source.id} - {source.snippet}"
                )
        lines.append("")
        lines.append("---")
        lines.append("")
    lines.append("INSIGNIFICANT VARIANCES")
    lines.extend(agent_run.insignificant or ["None"])
    return "\n".join(lines).strip() + "\n"


def write_markdown(commentary: str, dest_path: Path) -> None:
    dest_path = Path(dest_path)
    dest_path.parent.mkdir(parents=True, exist_ok=True)
    dest_path.write_text(commentary, encoding="utf-8", newline="\n")


def write_docx(commentary: str, dest_path: Path) -> None:
    dest_path = Path(dest_path)
    dest_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    lines = commentary.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal buffer
        body = "\n".join(buffer).strip()
        buffer = []
        if body:
            document.add_paragraph(body)

    for line in lines:
        stripped = line.strip()
        if stripped in _SECTION_TITLES:
            flush_paragraph()
            document.add_heading(stripped, level=1)
        elif not stripped:
            flush_paragraph()
        else:
            buffer.append(line)
    flush_paragraph()
    document.save(str(dest_path))


def export_from_run(agent_run: AgentRun, format: str, out_dir: Path) -> Path:
    out_dir = Path(out_dir)
    rendered = render_run_markdown(agent_run)
    ext = format.strip().lower().rstrip(".").lstrip(".")
    dest = out_dir / build_export_basename(agent_run.period_label, ext)
    if ext == "md":
        write_markdown(rendered, dest)
    elif ext == "docx":
        write_docx(rendered, dest)
    else:
        raise ValueError(f"Unsupported export format: {format}")
    return dest
